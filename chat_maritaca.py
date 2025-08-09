import os
import json
import openai
import PyPDF2
import docx
import streamlit as st
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import A4

MARITACA_API_KEY = st.secrets.get("MARITACA_API_KEY")

if not MARITACA_API_KEY:
    st.error("A variável MARITACA_API_KEY não foi definida em .streamlit/secrets.toml.")
    st.stop()

client = openai.OpenAI(
    api_key=MARITACA_API_KEY,
    base_url="https://chat.maritaca.ai/api"
)

def ler_pdf(arquivo):
    texto = ""
    leitor = PyPDF2.PdfReader(arquivo)
    for pagina in leitor.pages:
        texto_pagina = pagina.extract_text()
        if texto_pagina:
            texto += texto_pagina + "\n"
    return texto.strip()

def ler_txt(arquivo):
    return arquivo.read().decode("utf-8").strip()

def ler_docx(arquivo):
    doc = docx.Document(arquivo)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def criar_preprompt(texto):
    return {
        "role": "system",
        "content": (
            "Você é um assistente especialista em design educacional e criação de conteúdo. "
            "Considere o seguinte conteúdo base para todas as respostas:\n\n"
            f"{texto}"
        )
    }

def chat_with_bot(user_input, preprompt):
    try:
        response = client.chat.completions.create(
            model="sabiazim-3",
            messages=[preprompt, {"role": "user", "content": user_input}],
            temperature=0.7,
            max_tokens=2048
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Erro na comunicação com a API: {e}"

def gerar_glossario(conteudo, preprompt):
    prompt_glossario = """
    Seção 3. Glossário geral

    Com base no texto abaixo, identifique as palavras ou termos difíceis, técnicos ou pouco usuais para o público geral. 
    Para cada termo, forneça uma explicação clara e objetiva. 
    Apresente o glossário no seguinte formato, sem linhas em branco e sem repetições:

    N°:\tTermo:\tDefinição / significado:
    1\tpalavra\tdefinição
    2\tpalavra\tdefinição
    3\t...

    Texto base:
    """ + conteudo
    return chat_with_bot(prompt_glossario, preprompt)

def gerar_links_anexos(conteudo, preprompt):
    prompt_links = """
    Seção 4. Links de materiais complementares e anexos

    Analise o conteúdo abaixo, extraindo apenas os links e anexos citados no texto. 
    Organize-os em uma seção chamada "Links complementares e anexos", seguindo os exemplos de formatação abaixo. 
    Não invente links, use apenas os que realmente aparecem no texto.

    Exemplos de formato:

    Vídeos do youtube
    TÍTULO DA PARTE. Autor da postagem (duração do arquivo): idioma. Endereço disponível da web. Data de acesso.

    Blog
    TÍTULO DA PUBLICAÇÃO. Título da parte. Data de publicação. Endereço disponível da web. Data de acesso.

    Podcast
    SOBRENOME DO AUTOR, Nome. Título da parte. Duração da mídia. Data de publicação. Endereço disponível da web. Data de acesso.

    Conteúdo base:
    """ + conteudo
    return chat_with_bot(prompt_links, preprompt)

# --- STREAMLIT ---

st.title("Maritaca AI - Gerador de Módulo Educacional")

if "conteudo_modulo" not in st.session_state:
    st.session_state["conteudo_modulo"] = []
if "texto_final" not in st.session_state:
    st.session_state["texto_final"] = None

tema_geral = st.text_input("Digite uma breve descrição do tema geral:")
arquivo = st.file_uploader("Envie um arquivo (.pdf, .txt, .docx)", type=["pdf", "txt", "docx"])

if st.button("Processar") and tema_geral and arquivo:
    st.session_state["conteudo_modulo"] = []
    ext = arquivo.name.split(".")[-1].lower()
    if ext == "pdf":
        texto = ler_pdf(arquivo)
    elif ext == "txt":
        texto = ler_txt(arquivo)
    elif ext == "docx":
        texto = ler_docx(arquivo)
    else:
        st.error("Formato de arquivo não suportado.")
        st.stop()

    preprompt = criar_preprompt(f"Tema geral: {tema_geral}\n\n{texto}")

    # Seção 1. Introdução ao conteúdo
    prompt_introducao = """
    Seção 1. Introdução ao conteúdo

    Aja como um conteudista educacional experiente. Com base no conteúdo do sistema, redija um texto introdutório para um módulo educacional. Siga estas regras:

    - O texto deve ter **pelo menos 3 parágrafos**, totalizando **no máximo 1.200 caracteres sem espaço**.
    - Instigue o leitor a estudar o tema abordado.
    - Use o seguinte modelo estrutural:

    1. Parágrafo inicial com uma introdução envolvente sobre o tema geral do módulo.
    2. "Neste módulo, você vai aprender…" seguido de um resumo claro dos conteúdos abordados e sua importância para profissionais ou interessados na área.
    3. "Outro tema relevante é…" e depois "Por fim, você terá a oportunidade de…", encerrando com a ampliação de repertório do estudante.

    Finalize com: **Bons estudos!**
    """
    conteudo_introducao = chat_with_bot(prompt_introducao, preprompt)
    st.session_state["conteudo_modulo"].append("Seção 1. Introdução ao conteúdo\n" + conteudo_introducao)

    # Seção 2. Unidades de aprendizagem do Módulo
    prompt_unidades = """
    Seção 2. Unidades de aprendizagem do Módulo

    Atue como um especialista em design instrucional. Sua tarefa é desenvolver o conteúdo detalhado de um módulo educacional com base no conteúdo do sistema. Siga estritamente as diretrizes abaixo:

    1.  **Estrutura em Unidades:** Subdivida os temas principais do conteúdo em Unidades de aprendizagem. Dê um título claro e descritivo para cada Unidade (exemplo: "Unidade 1: A EaD e a Legislação Brasileira").
    
    2.  **Elaboração do Conteúdo:** Para cada unidade, elabore um texto didático que explique os conceitos essenciais do tema. O texto deve ser claro, objetivo e aprofundado.
    
    3.  **Fundamentação e Referências:** Fundamente os conceitos com referências (como livros, artigos, leis, ou outras fontes de credibilidade, preferencialmente disponíveis online). Cite as fontes diretamente no texto ou ao final de cada unidade.
    
    4.  **Sugestão de Elementos Visuais:** Onde for pertinente, sugira elementos visuais para enriquecer a aprendizagem. Use marcadores como `[Sugestão de Imagem: descrição da imagem]` ou `[Caixa de Destaque: texto a ser destacado]`.
    
    5.  **Restrições:**
        * Cada unidade deve ter, no máximo, **3.000 caracteres sem espaço**.
        * Ignore quaisquer menções a formatação de fonte (Arial, tamanho 12) ou termos específicos como "RAE". O foco é exclusivamente na qualidade e estrutura do conteúdo textual.

    Desenvolva a quantidade de unidades que julgar necessária para cobrir os tópicos centrais do material de base.
    """
    conteudo_unidades = chat_with_bot(prompt_unidades, preprompt)
    st.session_state["conteudo_modulo"].append("Seção 2. Unidades de aprendizagem do Módulo\n" + conteudo_unidades)

    # Seção 3. Glossário geral
    conteudo_para_glossario = f"{tema_geral}\n{texto}"
    conteudo_glossario = gerar_glossario(conteudo_para_glossario, preprompt)
    st.session_state["conteudo_modulo"].append("Seção 3. Glossário geral\n" + conteudo_glossario)

    # Seção 4. Links de materiais complementares e anexos
    conteudo_links = gerar_links_anexos(texto, preprompt)
    st.session_state["conteudo_modulo"].append("Seção 4. Links de materiais complementares e anexos\n" + conteudo_links)

    # Seção 5. Unidade de conclusão do módulo
    prompt_conclusao = """
    Seção 5. Unidade de conclusão do módulo

    Atue como um conteudista educacional experiente. Com base no conteúdo do sistema, escreva a unidade de conclusão do módulo, utilizando o mesmo formato da apresentação, mas agora recapitulando tudo o que a pessoa aprendeu. Siga estas orientações:

    - O texto deve ter pelo menos 3 parágrafos, totalizando no máximo 1.200 caracteres sem espaço.
    - Faça uma síntese dos principais pontos e temas abordados no módulo, destacando o que o estudante aprendeu.
    - Use frases como "Neste módulo, você aprendeu...", "Além disso, foi possível compreender...", "Por fim, você está apto a...".
    - Finalize com uma mensagem de incentivo ao estudante para aplicar o conhecimento adquirido.

    Finalize com: **Parabéns por concluir o módulo!**
    """
    conteudo_conclusao = chat_with_bot(prompt_conclusao, preprompt)
    st.session_state["conteudo_modulo"].append("Seção 5. Unidade de conclusão do módulo\n" + conteudo_conclusao)

    # Seção 6. Referências do Módulo
    prompt_referencias = """
    Seção 6. Referências do Módulo

    Com base no conteúdo do sistema, extraia e organize todas as referências bibliográficas citadas no texto do módulo. 
    Utilize o seguinte formato para cada referência:

    SOBRENOME, Nome do autor. Título. Edição. Local: Editora, data publicação.

    Liste apenas referências realmente presentes no conteúdo, sem inventar. 
    Caso não haja referências explícitas, retorne apenas: "Nenhuma referência encontrada no conteúdo."

    Apresente a lista de forma clara e sem repetições.
    """
    conteudo_referencias = chat_with_bot(prompt_referencias, preprompt)
    st.session_state["conteudo_modulo"].append("Seção 6. Referências do Módulo\n" + conteudo_referencias)

    texto_final = ""
    for secao in st.session_state.get("conteudo_modulo", []):
        texto_final += secao + "\n\n"
    st.session_state["texto_final"] = texto_final

# --- VISUALIZAÇÃO E BOTÕES DE DOWNLOAD ---
if st.session_state.get("texto_final"):
    st.markdown("---")
    texto_final = st.session_state["texto_final"]

    st.markdown("### Visualização do módulo gerado")
    st.markdown(f"<div style='white-space: pre-wrap'>{texto_final}</div>", unsafe_allow_html=True)

    # TXT
    buffer_txt = BytesIO()
    buffer_txt.write(texto_final.encode("utf-8"))
    buffer_txt.seek(0)
    st.download_button("📄 Baixar como TXT", buffer_txt, "modulo.txt", "text/plain", key="download-txt")

    # DOCX
    buffer_docx = BytesIO()
    doc = Document()
    for linha in texto_final.split("\n"):
        doc.add_paragraph(linha)
    doc.save(buffer_docx)
    buffer_docx.seek(0)
    st.download_button("📝 Baixar como DOCX", buffer_docx, "modulo.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="download-docx")

    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch

    buffer_pdf = BytesIO()
    doc_pdf = SimpleDocTemplate(
        buffer_pdf,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40,
    )
    styles = getSampleStyleSheet()
    story = []
    for linha in texto_final.split("\n"):
        if linha.strip() == "":
            story.append(Spacer(1, 0.2 * inch))
        else:
            story.append(Paragraph(linha, styles["Normal"]))
    doc_pdf.build(story)
    buffer_pdf.seek(0)
    st.download_button("📕 Baixar como PDF", buffer_pdf, "modulo.pdf", "application/pdf", key="download-pdf")

    st.markdown(
        "<div style='position: fixed; bottom: 8px; right: 16px; font-size: 10px; color: #888;'>"
        "Feito por: PietroTy, 2025"
        "</div>",
        unsafe_allow_html=True
    )