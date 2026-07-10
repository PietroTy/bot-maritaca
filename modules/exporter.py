"""
modules/exporter.py — Módulo 5: Exportação de Documentos
Escriba

ESTADO ATUAL: Exportação funcional para PDF (ReportLab) e TXT.
ROADMAP: 
  - LaTeX: Gerar .tex via template e compilar para PDF com pdflatex
  - DOCX: Estilos ABNT corretos (margens, fontes, espaçamento)
  - Integração com Overleaf via API

Responsabilidades:
- Montar o documento final com cabeçalho, seções e rodapé
- Gerar PDF (ReportLab), TXT e DOCX (python-docx)
- Retornar bytes prontos para download no Streamlit
"""

import io
import re
from datetime import datetime
from typing import Optional


# ─────────────────────────────────────────────────────────────
# Filtro de Sanitização Pós-Geração
# Remove artefatos do sistema que o LLM possa ter inserido
# apesar das regras do system prompt.
# ─────────────────────────────────────────────────────────────
_PADROES_LIXO = [
    # Assinatura do sistema no corpo do texto
    r"Escriba\s+v\d+\.\d+[^\n]*",
    # Tags de verificação residuais: [LEI 15.388/2026 VERIFICAR FONTE], [VERIFICAR], etc.
    r"\[\s*[^\]]*?VERIFICAR[^\]]*?\]",
    # Tags genéricas de marcação interna: [NOTA], [REVISAR], [CHECAR FONTE]
    r"\[\s*(?:NOTA|REVISAR|CHECAR FONTE|FONTE VERIFICAR|CONFIRMAR)[^\]]*?\]",
    # Marcadores de página soltos: "Página 11", "Pág. 11"
    r"(?:P[áa]gina|P[áa]g\.?)\s+\d+",
]
_RE_SANITIZAR = re.compile("|".join(_PADROES_LIXO), re.IGNORECASE)


def _sanitizar_texto(texto: str) -> str:
    """Remove metadados, tags de sistema e assinaturas do texto gerado pela IA."""
    # Remove os padrões conhecidos
    texto = _RE_SANITIZAR.sub("", texto)
    # Normaliza espaços excessivos deixados pelas remoções
    texto = re.sub(r" {2,}", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


# --- PDF via ReportLab ---
def _gerar_pdf(secoes: list, tema: str, idioma: str) -> bytes:
    """Gera PDF formatado com ReportLab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    except ImportError:
        raise ImportError("reportlab não instalado. Execute: pip install reportlab")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=48, leftMargin=48, topMargin=56, bottomMargin=56,
    )
    base_font = "Helvetica"
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle("TitleMain", parent=styles["Title"],
        fontName=base_font, fontSize=20, leading=24, alignment=TA_CENTER, spaceAfter=18))
    styles.add(ParagraphStyle("ModuleMeta", parent=styles["Normal"],
        fontName=base_font, fontSize=10, leading=12, alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle("HeadingSection", parent=styles["Heading2"],
        fontName=base_font, fontSize=14, leading=18, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle("EscribaBody", parent=styles["Normal"],
        fontName=base_font, fontSize=11, leading=15, spaceBefore=6, spaceAfter=6, alignment=TA_LEFT))

    def draw_page(canvas, doc):
        canvas.saveState()
        w, h = A4
        canvas.setFont(base_font, 8)
        canvas.setFillColorRGB(0.4, 0.4, 0.4)
        canvas.drawCentredString(w / 2.0, 20,
            f"Escriba — {tema}    •    Página {doc.page}")
        canvas.restoreState()

    story = []

    if tema:
        story.append(Paragraph(tema, styles["HeadingSection"]))
        story.append(Spacer(1, 0.2 * inch))

    import re

    for secao in secoes:
        titulo = secao.get("titulo", "")
        texto = secao.get("texto", "")
        if titulo:
            story.append(Paragraph(titulo, styles["HeadingSection"]))

        texto = _sanitizar_texto(texto)
        for paragrafo in [p.strip() for p in texto.split("\n\n") if p.strip()]:
            # Convert Markdown headers
            header_match = re.match(r'^(#{1,6})\s+(.*)', paragrafo)
            if header_match:
                story.append(Paragraph(header_match.group(2).strip(), styles["HeadingSection"]))
                continue

            p_formatado = paragrafo.replace("\n", "<br/>")
            # Convert bold and italic for ReportLab HTML-like tags
            p_formatado = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', p_formatado)
            p_formatado = re.sub(r'\*([^\s\*].*?[^\s\*])\*', r'<i>\1</i>', p_formatado)

            story.append(Paragraph(p_formatado, styles["EscribaBody"]))

        story.append(Spacer(1, 0.12 * inch))

    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    buffer.seek(0)
    return buffer.read()


# --- TXT simples ---
def _gerar_txt(secoes: list, tema: str, idioma: str) -> bytes:
    """Gera arquivo TXT simples."""
    linhas = []
    if tema:
        linhas.append(f"TEMA: {tema}\n{'=' * 60}\n")
    
    for secao in secoes:
        if secao.get("titulo", ""):
            linhas.append(f"\n{'=' * 40}")
            linhas.append(secao.get("titulo", ""))
            linhas.append("=" * 40)
        linhas.append(_sanitizar_texto(secao.get("texto", "")))
    return "\n".join(linhas).encode("utf-8")


# --- DOCX via python-docx ---
def _gerar_docx(secoes: list, tema: str, idioma: str) -> bytes:
    """Gera DOCX. ROADMAP: Aplicar estilos ABNT corretos."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
    except ImportError:
        raise ImportError("python-docx não instalado. Execute: pip install python-docx")

    doc = DocxDocument()
    if tema:
        doc.add_heading(tema, level=0)

    import re

    for secao in secoes:
        if secao.get("titulo", ""):
            doc.add_heading(secao.get("titulo", ""), level=2)
        texto = _sanitizar_texto(secao.get("texto", ""))

        for paragrafo in [p.strip() for p in texto.split("\n\n") if p.strip()]:
            # Suporte para Cabeçalhos Markdown
            header_match = re.match(r'^(#{1,6})\s+(.*)', paragrafo)
            if header_match:
                level = len(header_match.group(1))
                doc.add_heading(header_match.group(2).strip(), level=level)
                continue

            p_obj = doc.add_paragraph()
            # Suporte para Negrito e Itálico simultâneos
            parts = re.split(r'(\*\*.*?\*\*|\*[^\s\*].*?[^\s\*]\*)', paragrafo)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p_obj.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    run = p_obj.add_run(part[1:-1])
                    run.italic = True
                else:
                    p_obj.add_run(part)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# --- [ROADMAP] LaTeX ---
def _gerar_tex(secoes: list, tema: str, idioma: str) -> bytes:
    """
    [ROADMAP] Gera arquivo .tex a partir do template LaTeX.
    Futuramente: compilar com pdflatex localmente ou via Docker.
    """
    linhas = [
        r"\documentclass[12pt,a4paper]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage[brazil]{babel}",
        r"\usepackage{geometry}",
        r"\geometry{a4paper, margin=2.5cm}",
        r"\begin{document}",
        r"\begin{center}",
        rf"\textbf{{\LARGE Escriba}}\\[0.5em]",
        rf"\textbf{{\large {tema}}}\\[0.3em]",
        rf"\small Idioma: {idioma} | Gerado: {datetime.utcnow().strftime('%Y-%m-%d')}",
        r"\end{center}",
        r"\newpage",
    ]
    for secao in secoes:
        titulo = secao.get("titulo", "").replace("_", r"\_")
        texto = secao.get("texto", "").replace("_", r"\_").replace("&", r"\&").replace("%", r"\%")
        linhas.append(rf"\section{{{titulo}}}")
        linhas.append(texto)
        linhas.append("")
    linhas.append(r"\end{document}")
    return "\n".join(linhas).encode("utf-8")


def export(
    polish_results: list,
    formato: str,
    tema: str,
    idioma: str,
    status_callback=None,
) -> tuple[bytes, str, str]:
    """
    Ponto de entrada principal do Exporter.

    Args:
        polish_results: Lista de PolishResult do Polisher.
        formato: "pdf", "txt", "docx" ou "tex".
        tema: Tema geral para metadados do documento.
        idioma: Idioma de saída.
        status_callback: Função para mensagens de status na UI.

    Returns:
        Tupla (bytes_do_arquivo, nome_arquivo, mime_type).
    """
    if status_callback:
        status_callback(f"Compilando documento final em formato {formato.upper()}...")

    secoes = [
        {"titulo": r.secao_id.replace("_", " ").title(), "texto": r.texto_polido}
        for r in polish_results
    ]
    # Usa titulo correto se disponível
    for r, s in zip(polish_results, secoes):
        # Tenta pegar titulo do objeto GeneratorResult original
        if hasattr(r, 'secao_titulo') and r.secao_titulo:
            s["titulo"] = r.secao_titulo

    formato = formato.lower()
    if formato == "pdf":
        conteudo = _gerar_pdf(secoes, tema, idioma)
        nome = "modulo_escriba.pdf"
        mime = "application/pdf"
    elif formato == "txt":
        conteudo = _gerar_txt(secoes, tema, idioma)
        nome = "modulo_escriba.txt"
        mime = "text/plain"
    elif formato == "docx":
        conteudo = _gerar_docx(secoes, tema, idioma)
        nome = "modulo_escriba.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif formato == "tex":
        conteudo = _gerar_tex(secoes, tema, idioma)
        nome = "modulo_escriba.tex"
        mime = "text/x-tex"
    else:
        raise ValueError(f"Formato '{formato}' não suportado.")

    if status_callback:
        status_callback(f"✅ Exportação concluída: {nome} ({len(conteudo):,} bytes).")

    return conteudo, nome, mime
